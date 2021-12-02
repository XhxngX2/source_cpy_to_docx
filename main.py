import os, argparse
from docx import Document

C_CPP_FILE_ENDINGS = [".c", ".h", ".cpp", ".asm"]
PYTHON_FILE_ENDINGS = [".py"]


def specific_end_file2docxfile(out_docx_document, file_path, allowed_ends):
    for end in allowed_ends:
        if file_path.endswith(end):
            file_name_with_end = file_path.split("/")[-1]
            out_docx_document.add_heading("\n" + file_name_with_end, 1)
            with open(file_path) as fp:
                try:
                    content = fp.read()
                    out_docx_document.add_paragraph(content)
                except Exception as e:
                    print("Error occur during reading " + file_name_with_end+":")
                    print(str(e))


def file_operation(out_docx_document, file_path, project_type):
    if project_type == 'C/C++':
        specific_end_file2docxfile(out_docx_document, file_path, C_CPP_FILE_ENDINGS)

    if project_type == 'Python':
        specific_end_file2docxfile(out_docx_document, file_path, PYTHON_FILE_ENDINGS)


def convert(working_dir, out_docx, proj_type):
    dir_contents = os.listdir(working_dir)
    content_num = len(dir_contents)

    if content_num == 0:
        print(working_dir + 'is an empty folder')
    else:
        for item_index in range(content_num):
            item = dir_contents[item_index]
            item_path = os.path.join(working_dir, item)

            if os.path.isdir(item_path):
                convert(item_path, out_docx, proj_type)

            if os.path.isfile(item_path):
                file_operation(out_docx_document=out_docx, file_path=item_path, project_type=proj_type)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(prog="\n\rIterate through folders to paste source code into FOLDERNAME.docx",
                                     epilog="\n\rxhxngx2")

    parser.add_argument('-d', '--dir', help="\n\rSpecify the Source code dir.",
                        required=True)
    parser.add_argument('-f', '--file', choices=['C/C++', 'Python'],
                        required=True, help="\n\rSpecify the Project kind.")
    args = parser.parse_args()
    proj_name = args.dir.split("/")[-1]

    # create a docx file
    docx_obj = Document()
    docx_obj.add_heading(proj_name, 0)

    convert(working_dir=args.dir, out_docx=docx_obj, proj_type=args.file)

    file_name = proj_name  # in case we need some additional descriptions
    docx_obj.save(file_name + ".docx")
